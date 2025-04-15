import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#Ruta del archivo excel
ruta_excel=r'H:\Proyectos\Python\Prueba Isat.xlsx'

#Leer sin encabezados
df=pd.read_excel(ruta_excel,header=None,engine='openpyxl')

if len(df) > 0:
    fila=df.iloc[0]
    doc=Document()

    for valor in fila:
        p=doc.add_paragraph()
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        run=p.add_run(str(valor))
        run.font.size=Pt(48)
        doc.add_page_break()

    doc.save('Checa.docx')
    print("Documento generado correctamente.")
else:
    print("El archivo no tiene datos.")