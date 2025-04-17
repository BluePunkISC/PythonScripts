from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#Ruta del archivo
ruta_excel=r'H:\Proyectos\Python\Prueba1.xlsm'

#Cargar excel
wb=load_workbook(ruta_excel, data_only=True)
ws=wb.active

#Crear word
doc=Document()

#Parte 1: Bloque de celdas combinadas
for fila in range(13,88,5):
    codigo=ws[f'B{fila}'].value
    partes_texto=[ws[f'G{i}'].value for i in range(fila, fila+5)]
    partes_texto=[str(p).strip() for p in partes_texto if p]

    if codigo and partes_texto:
        texto_final=f"{codigo}: {' - '.join(partes_texto)}"
        p=doc.add_paragraph()
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        run=p.add_run(texto_final)
        run.font.size=Pt(48)
        doc.add_page_break()

#Parte 2: Columnas combinadas de A a U
filas_horizontal=[88,100,109,124,136]
for fila in filas_horizontal:
    texto_horizontal=[]
    for col in range(1,22): #Columnas A (1) U (21)
        val=ws.cell(row=fila, column=col).value
        if val:
            texto_horizontal.append(str(val).strip())

    if texto_horizontal:
        linea=' '.join(texto_horizontal)
        p=doc.add_paragraph()
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        run=p.add_run(linea)
        run.font.size=Pt(36)
        doc.add_page_break()

#Parte 3: Filas sueltas de B + G
filas_sueltas=list(range(89,100)) + list(range(101,109)) + list(range(110,124)) + list(range(137,139))
for fila in filas_sueltas:
    codigo=ws[f'B{fila}'].value
    texto=ws[f'G{fila}'].value

    if codigo and texto:
        texto_final=f'{codigo}: {str(texto).strip()}'
        p=doc.add_paragraph()
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        run=p.add_run(texto_final)
        run.font.size=Pt(48)
        doc.add_page_break()

#Guardar documento
doc.save("Concatenado.docx")
print("Documento generado con éxito.")